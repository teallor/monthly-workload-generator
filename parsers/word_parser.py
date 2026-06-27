from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from rules import hours_from_time
from .common import CourseRecord


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _split_time(value: str) -> tuple[str, str]:
    found = re.findall(r"\d{1,2}:\d{2}", value)
    return (found[0], found[1]) if len(found) >= 2 else (value, "")


def _date(value: str, year: int) -> str:
    match = re.search(r"(\d{1,2})\s*月\s*(\d{1,2})\s*日", value)
    if not match:
        return ""
    return f"{year:04d}-{int(match.group(1)):02d}-{int(match.group(2)):02d}"


def _unique_header_columns(table) -> dict[str, int]:
    result: dict[str, int] = {}
    for index, cell in enumerate(table.rows[0].cells):
        label = _clean(cell.text)
        if label and label not in result:
            result[label] = index
    return result


def _parse_schedule_table(table, source: Path, year: int, context: dict) -> list[CourseRecord]:
    header_row = None
    for i, row in enumerate(table.rows):
        labels = [_clean(cell.text) for cell in row.cells]
        if any("日期" in x for x in labels) and any("课程" in x for x in labels):
            header_row = i
            break
    if header_row is None:
        return []
    headers = _unique_header_columns(type("T", (), {"rows": table.rows[header_row:]})())

    def col(*keys: str):
        for label, index in headers.items():
            if any(key in label for key in keys):
                return index
        return None

    date_col, time_col = col("日期"), col("时间")
    course_col, teacher_col = col("课程", "项目"), col("培训师", "授课老师", "教师")
    hours_col = col("课时")
    if None in (date_col, time_col, course_col, teacher_col):
        return []
    records: list[CourseRecord] = []
    inherited_date = ""
    for row in table.rows[header_row + 1:]:
        cells = [_clean(cell.text) for cell in row.cells]
        raw_date = cells[date_col]
        if _date(raw_date, year):
            inherited_date = _date(raw_date, year)
        teacher = cells[teacher_col]
        course = cells[course_col]
        if not teacher or not course:
            continue
        time_text = cells[time_col]
        start, end = _split_time(time_text)
        explicit = cells[hours_col] if hours_col is not None else ""
        if not re.search(r"\d{1,2}:\d{2}|全天", time_text) and not re.fullmatch(r"\d+(?:\.\d+)?", explicit or ""):
            continue
        try:
            hours = float(explicit)
            if hours.is_integer():
                hours = int(hours)
        except (TypeError, ValueError):
            hours = hours_from_time(time_text)
        records.append(CourseRecord(
            source_file=source.name, date=inherited_date, start_time=start, end_time=end,
            course_name=course, teacher=teacher, hours=hours,
            project=context.get("project", source.stem), audience=context.get("audience", ""),
            location=context.get("location", ""), context=" | ".join(cells),
            confidence=0.98, is_delivery=context.get("is_delivery", False),
        ))
    return records


def parse_docx(path: Path, year: int) -> list[CourseRecord]:
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            full_text += "\n" + " | ".join(_clean(c.text) for c in row.cells)
    years = re.findall(r"(20\d{2})\s*年", full_text + " " + path.name)
    document_year = int(years[0]) if years else year
    context = {"project": path.stem, "audience": "", "location": "", "is_delivery": False}
    if "工作效能提速训练营" in full_text:
        context.update(project="2026年上港集团工作效能提速训练营", audience="集团各单位")
        match = re.search(r"培训地点[：:]?\s*([^\n]+)", full_text)
        if match:
            context["location"] = _clean(match.group(1))
    if "送教上门" in full_text:
        context["is_delivery"] = True
        unit = "九江公司" if "九江公司" in full_text else ""
        context.update(project=f"{unit}送教上门培训" if unit else "送教上门培训", audience=unit, location=unit)
    records: list[CourseRecord] = []
    for table in doc.tables:
        records.extend(_parse_schedule_table(table, path, document_year, context))
    return records


def parse_doc(path: Path, year: int) -> list[CourseRecord]:
    try:
        import win32com.client
    except ImportError:
        win32com = None
    with tempfile.TemporaryDirectory(prefix="workload_doc_") as folder:
        converted = Path(folder) / f"{path.stem}.docx"
        if win32com is not None:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
            try:
                doc.SaveAs2(str(converted), FileFormat=16)
            finally:
                doc.Close(False)
                word.Quit()
        else:
            bridge = Path(__file__).resolve().parents[1] / "office_bridge.ps1"
            subprocess.run([
                "powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(bridge),
                "-Action", "convert-doc", "-Path", str(path.resolve()), "-Output", str(converted),
            ], check=True, capture_output=True, text=True, encoding="utf-8", errors="replace")
        records = parse_docx(converted, year)
        for record in records:
            record.source_file = path.name
        return records
